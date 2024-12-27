from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QMessageBox, QTableWidget, QTableWidgetItem, QFileDialog
)
from database import get_connection

try:
    import docx
    from docx import Document
except ImportError:
    Document = None

class DocInvoicesForm(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Документы: Накладные (WarehouseInvoices)")
        layout_main = QVBoxLayout()

        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Номер", "Дата", "Поставщик ID", "Сумма"])
        layout_main.addWidget(self.table)

        btn_export = QPushButton("Вывести в Word")
        btn_export.clicked.connect(self.export_to_word)
        layout_main.addWidget(btn_export)

        self.setLayout(layout_main)
        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, invoice_number, date, supplier_id, total_sum
            FROM WarehouseInvoices
            ORDER BY id ASC
        """)
        rows = cursor.fetchall()
        self.table.setRowCount(len(rows))
        for i, (inv_id, inv_num, inv_date, sup_id, t_sum) in enumerate(rows):
            self.table.setItem(i, 0, QTableWidgetItem(str(inv_id)))
            self.table.setItem(i, 1, QTableWidgetItem(str(inv_num)))
            self.table.setItem(i, 2, QTableWidgetItem(str(inv_date)))
            self.table.setItem(i, 3, QTableWidgetItem(str(sup_id)))
            self.table.setItem(i, 4, QTableWidgetItem(str(t_sum) if t_sum else "0"))
        conn.close()

    def export_to_word(self):
        if Document is None:
            QMessageBox.warning(self, "Ошибка", "Модуль 'python-docx' не установлен.")
            return
        row_count = self.table.rowCount()
        if row_count == 0:
            QMessageBox.information(self, "Нет данных", "Нет накладных для вывода.")
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить накладные в Word", "", "Word Document (*.docx)")
        if not file_path:
            return

        document = Document()
        document.add_heading("Накладные (WarehouseInvoices)", 0)

        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Номер"
        hdr_cells[2].text = "Дата"
        hdr_cells[3].text = "Поставщик"
        hdr_cells[4].text = "Сумма"

        for i in range(row_count):
            inv_id = self.table.item(i, 0).text()
            inv_num = self.table.item(i, 1).text()
            inv_date = self.table.item(i, 2).text()
            sup_id = self.table.item(i, 3).text()
            t_sum = self.table.item(i, 4).text()

            row_cells = table.add_row().cells
            row_cells[0].text = inv_id
            row_cells[1].text = inv_num
            row_cells[2].text = inv_date
            row_cells[3].text = sup_id
            row_cells[4].text = t_sum

        document.save(file_path)
        QMessageBox.information(self, "Ок", f"Файл Word сохранён: {file_path}")
