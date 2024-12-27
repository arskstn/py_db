
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QMessageBox, QTableWidget, QTableWidgetItem, QFileDialog
)
from database import get_connection

try:
    import docx
    from docx import Document
except ImportError:
    Document = None

class DocOrdersForm(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Документы: Заказы (WarehouseOrders)")

        layout_main = QVBoxLayout()

        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Номер заказа", "Дата", "Клиент ID", "Сумма"])
        layout_main.addWidget(self.table)

        btn_export = QPushButton("Вывести в Word")
        btn_export.clicked.connect(self.export_to_word)
        layout_main.addWidget(btn_export)

        self.setLayout(layout_main)
        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, order_number, date, client_id, total_sum
            FROM WarehouseOrders
            ORDER BY id ASC
        """)
        rows = cursor.fetchall()
        self.table.setRowCount(len(rows))
        for i, (oid, onum, odate, cid, t_sum) in enumerate(rows):
            self.table.setItem(i, 0, QTableWidgetItem(str(oid)))
            self.table.setItem(i, 1, QTableWidgetItem(str(onum)))
            self.table.setItem(i, 2, QTableWidgetItem(str(odate)))
            self.table.setItem(i, 3, QTableWidgetItem(str(cid)))
            self.table.setItem(i, 4, QTableWidgetItem(str(t_sum) if t_sum else "0"))
        conn.close()

    def export_to_word(self):
        if Document is None:
            QMessageBox.warning(self, "Ошибка", "Модуль 'python-docx' не установлен.")
            return
        row_count = self.table.rowCount()
        if row_count == 0:
            QMessageBox.information(self, "Нет данных", "Нет заказов для вывода.")
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить заказы в Word", "", "Word Document (*.docx)")
        if not file_path:
            return

        document = Document()
        document.add_heading("Заказы (WarehouseOrders)", 0)

        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Номер"
        hdr_cells[2].text = "Дата"
        hdr_cells[3].text = "Клиент ID"
        hdr_cells[4].text = "Сумма"

        for i in range(row_count):
            o_id = self.table.item(i, 0).text()
            o_num = self.table.item(i, 1).text()
            o_date = self.table.item(i, 2).text()
            c_id = self.table.item(i, 3).text()
            t_sum = self.table.item(i, 4).text()

            row_cells = table.add_row().cells
            row_cells[0].text = o_id
            row_cells[1].text = o_num
            row_cells[2].text = o_date
            row_cells[3].text = c_id
            row_cells[4].text = t_sum

        document.save(file_path)
        QMessageBox.information(self, "Ок", f"Файл Word сохранён: {file_path}")
